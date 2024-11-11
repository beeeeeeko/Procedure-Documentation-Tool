import os
import glob
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import messagebox

def create_instruction_docx(text_file, image_dir, output_file):
    if not os.path.isfile(text_file):
        print("指定されたテキストファイルは存在しません。")
        return
    
    if not os.path.isdir(image_dir):
        print("指定された画像ディレクトリが存在しません。")
        return

    try:
        with open(text_file, 'r', encoding='utf-8') as file:
            instructions = file.readlines()
    except Exception as e:
        print(f"テキストファイルの読み込みに失敗しました: {e}")
        return
    

    doc = Document()

    # タイトル設定
    title = extract_title(instructions)
    add_title_page(doc, title)

    # 画像取得
    images = gather_images(image_dir)

    # ドキュメント作成
    add_instructions(doc, instructions, images)

    try:
        doc.save(output_file)
        print(f"手順書が正常に保存されました: {output_file}")
    except Exception as e:
        print(f"ファイルの保存に失敗しました: {e}")

# ドキュメントの保存に成功した場合
    try:
        doc.save(output_file)
        messagebox.showinfo("成功", f"手順書が正常に保存されました: {output_file}")
    except Exception as e:
        messagebox.showerror("エラー", f"ファイルの保存に失敗しました: {e}")

def extract_title(lines):
    """タイトル取得用のヘルパー関数"""
    if lines and lines[0].strip().startswith('"') and lines[0].strip().endswith('"'):
        return lines.pop(0).strip().strip('"')
    return "Untitled Document"

def add_title_page(doc, title):
    """タイトルページ作成のためのヘルパー関数"""
    title_paragraph = doc.add_paragraph()
    run = title_paragraph.add_run(title)
    run.bold = True
    run.font.size = Inches(1)  # サイズをインチで設定

    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_paragraph.space_before = Inches(2)  # ページ中央に配置

    doc.add_page_break()

def gather_images(directory):
    """画像収集のためのヘルパー関数"""
    return sorted(glob.glob(os.path.join(directory, '*.jpg'))) + sorted(glob.glob(os.path.join(directory, '*.png')))

def add_instructions(doc, instructions, images):
    """手順書コンテンツを追加"""
    heading_counter = subheading_counter = paragraph_counter = image_index = 1

    for line in instructions:
        line = line.strip()
        if line.startswith('#'):
            level = line.count('#')
            if level == 1:
                doc.add_heading(f"{heading_counter}. {line.strip('# ')}", level=1)
                heading_counter += 1
                subheading_counter = 1
            else:
                doc.add_heading(f"{heading_counter - 1}.{subheading_counter} {line.strip('# ')}", level=level)
                subheading_counter += 1
            paragraph_counter = 1
        else:
            if line:
                doc.add_paragraph(f"({paragraph_counter}) {line}")
                paragraph_counter += 1
                if image_index < len(images):
                    doc.add_picture(images[image_index], width=Inches(3))
                    image_index += 1
